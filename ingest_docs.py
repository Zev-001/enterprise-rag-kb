"""
ingest_docs.py —— 把上传的公司文档解析成纯文本，再喂给 rag_core 切块入库。

支持格式：
    .pdf    pdfplumber（能读文字层的 PDF；扫描件图片 PDF 读不出字，会提示）
    .docx   python-docx（Word 2007+；老 .doc 不支持）
    .txt    UTF-8 纯文本
    .md     Markdown，先去标记符号再存（保留可读文字）

设计取舍（给小白看的）：
    1. 只存「原文文字」，不存摘要。摘要已经过模型加工，再把摘要当资料检索是二级转述。
       存原文，溯源能追到具体哪份文件哪一段。
    2. 每块带 doc_id / filename / title 元数据，回答时能说「这句话来自《员工手册》第 3 段」。
    3. 大文件截断到 5 万字上限，避免一次塞太多把上下文撑爆（企业真实文档常很长）。
"""

import os

from rag_core import ingest_documents, kb_stats

MAX_CHARS = 50_000  # 单份文档最多摄入字数，超出截断


def _read_pdf(path):
    """读 PDF：文字层 + 表格层 + 扫描件 OCR 兜底（D3 补强）。

    - 文字层：page.extract_text()，常规 PDF 走这里。
    - 表格层：制度/手册的表格常含关键信息（报销标准、职级对照），用
      page.extract_tables() 抽出来单独成块，避免被当成普通段落漏检。
    - 扫描件：文字层为空的页，尝试 OCR（需装 pytesseract + poppler）。
      没装 OCR 库时优雅降级——记录提示，绝不打断整本入库。
    """
    import pdfplumber
    parts = []
    scanned_pages = 0
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(txt)
            else:
                # 文字层为空 → 可能是扫描件图片，尝试 OCR
                ocr_txt = _ocr_page(page)
                if ocr_txt and ocr_txt.strip():
                    parts.append(ocr_txt)
                else:
                    scanned_pages += 1
            # 表格层：抽出来单列成「【表格】」块
            try:
                for tbl in page.extract_tables():
                    rows = [" | ".join(c.strip() for c in row if c and c.strip())
                            for row in tbl if any(c and c.strip() for c in row)]
                    if rows:
                        parts.append("【表格】\n" + "\n".join(rows))
            except Exception:
                pass
    if scanned_pages:
        parts.append(
            "（注：本 PDF 有 {0} 页无文字层、疑似扫描件；未安装 OCR 库"
            "（pytesseract + poppler）时无法识别文字，装好后会自动转文字入库。）"
            .format(scanned_pages)
        )
    return "\n\n".join(parts)


def _ocr_page(page):
    """对单页做 OCR；缺依赖或失败一律返回 None（绝不打断入库）。"""
    try:
        import pytesseract  # 仅当装了 OCR 才走这条路径
    except Exception:
        return None
    try:
        img = page.to_image(resolution=200).original
        return pytesseract.image_to_string(img, lang="chi_sim+eng")
    except Exception:
        return None


def _read_docx(path):
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # 表格也抓出来，制度/手册常放表格里
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _read_md(path):
    import re
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    # 去掉 Markdown 标记，保留可读文字
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.M)          # 标题 #
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)                # 粗体
    text = re.sub(r"\*(.*?)\*", r"\1", text)                    # 斜体
    text = re.sub(r"`{1,3}(.*?)`{1,3}", r"\1", text, flags=re.S)  # 行内/代码块
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)                 # 图片
    text = re.sub(r"\[([^\]]+)\]\(.*?\)", r"\1", text)          # 链接保留文字
    text = re.sub(r"^\s*[-*+]\s+", "- ", text, flags=re.M)      # 列表符
    return text


_READERS = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".txt": _read_txt,
    ".md": _read_md,
}


def parse_file(path):
    """解析单个文件，返回 (filename, text) 或抛错。"""
    ext = os.path.splitext(path)[1].lower()
    if ext not in _READERS:
        raise ValueError("不支持的格式：{0}（支持 pdf/docx/txt/md）".format(ext))
    text = _READERS[ext](path)
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]
    return os.path.basename(path), text.strip()


def ingest_file(path, namespace="default", title=None):
    """解析并入库单个文件，返回 {filename, chars, chunks_added}。"""
    filename, text = parse_file(path)
    if not text:
        raise ValueError("从 {0} 没提取到任何文字（可能是扫描件 PDF，需要 OCR）".format(filename))
    docs = [{"filename": filename, "title": title or filename, "text": text}]
    added = ingest_documents(docs, namespace=namespace)
    return {"filename": filename, "chars": len(text), "chunks_added": added}


def ingest_paths(paths, namespace="default"):
    """批量入库，返回结果列表（含错误）。"""
    results = []
    for p in paths:
        try:
            r = ingest_file(p, namespace=namespace)
            results.append({"ok": True, **r})
        except Exception as e:
            results.append({"ok": False, "filename": os.path.basename(p), "error": str(e)})
    return results


# ----------------------------------------------------------------- CLI
def _main():
    import sys
    import glob as _glob
    args = sys.argv[1:]
    if not args:
        print("用法：python ingest_docs.py 文件1.pdf 文件2.docx ...")
        return
    paths = []
    for a in args:
        paths.extend(_glob.glob(a)) if any(c in a for c in "*?") else paths.append(a)
    res = ingest_paths(paths)
    for r in res:
        if r["ok"]:
            print("✅ {0}：{1} 字 → 入库 {2} 块".format(r["filename"], r["chars"], r["chunks_added"]))
        else:
            print("❌ {0}：{1}".format(r["filename"], r["error"]))
    st = kb_stats()
    print("--- 当前知识库：{0} 份文档，{1} 块 ---".format(len(st["docs"]), st["chunks"]))


if __name__ == "__main__":
    _main()
